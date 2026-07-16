"""
文件名: app/services/resume_service.py
创建时间: 2026-06-26
作者: TalentSense Team
功能描述: 简历服务，对应 API-Design.md 二、Resumes
入参: 文件字节 / resume_id / 查询条件
出参: UploadResponse / ResumeDetail / PageResult
对应 Business-Requirements F02-F06
"""
import json
import re
import uuid
from datetime import datetime, timezone
from app.core.config import settings
from app.core.exceptions import NotFoundError, ResumeParseError
from app.core.logger import logger
from app.core.minio_client import minio_client
from app.core.ocr import ocr_engine
from app.core.llm_client import llm_client
from app.core.embedding import embedding_model
from app.core.vector_store import vector_store
from app.core.database import MongoDB
from app.utils.pii import mask_phone, mask_email, hash_phone, hash_email
from app.utils.dedup import DedupChecker
from app.utils.chunker import split_parent_child
from app.utils.salary import parse_salary


class ResumeService:
    """简历服务"""

    def __init__(self):
        self.minio = minio_client
        self.ocr = ocr_engine
        self.llm = llm_client
        self.embedding = embedding_model
        self.vector_store = vector_store
        self.dedup_checker = None

    @property
    def resumes_coll(self):
        """延迟获取 MongoDB resumes collection（避免模块导入时 MongoDB 未连接）"""
        if hasattr(self, "_resumes_coll"):
            return self._resumes_coll
        return MongoDB.db.resumes if MongoDB.db is not None else None

    @resumes_coll.setter
    def resumes_coll(self, value):
        """测试注入用"""
        self._resumes_coll = value

    async def upload(self, file_bytes: bytes, file_name: str, content_type: str, overwrite: bool = False, background_tasks=None) -> dict:
        """AC2.1: 上传文件 → MinIO → 后台异步解析

        入参:
            file_bytes: 文件字节
            file_name: 文件名
            content_type: MIME 类型
            overwrite: 重复时是否覆盖
            background_tasks: FastAPI BackgroundTasks 实例，传入则后台解析；否则同步解析（兼容旧测试）
        出参:
            {"resume_id", "candidate_id", "file_name", "parse_status", "is_duplicate", "duplicate_with"}
        """
        file_id = self.minio.upload_bytes(file_bytes, file_name, content_type)
        resume_id = f"res_{uuid.uuid4().hex[:12]}"
        candidate_id = f"cand_{uuid.uuid4().hex[:12]}"
        now = datetime.now(timezone.utc).isoformat()
        await self.resumes_coll.insert_one({
            "resume_id": resume_id, "candidate_id": candidate_id,
            "file_info": {"file_id": file_id, "file_name": file_name, "file_type": file_name.split(".")[-1]},
            "parse_info": {"parse_status": "parsing", "parse_time": None},
            "tags": [], "is_favorite": False, "notes": "",
            "created_at": now, "updated_at": now,
        })
        # 解析较慢（LLM + BGE-M3），通过 BackgroundTasks 后台执行
        # 避免前端 HTTP 请求超时；前端通过列表刷新查看最终 parse_status
        if background_tasks is not None:
            background_tasks.add_task(self._parse_and_index_safe, resume_id, file_bytes, file_id, file_name, overwrite)
        else:
            await self._parse_and_index(resume_id, file_bytes, file_id, file_name, overwrite)
        return {
            "resume_id": resume_id, "candidate_id": candidate_id, "file_name": file_name,
            "parse_status": "parsing", "is_duplicate": False, "duplicate_with": None,
        }

    async def _parse_and_index_safe(self, resume_id: str, file_bytes: bytes, file_id: str, file_name: str, overwrite: bool):
        """BackgroundTasks 包装：同步等待 _parse_and_index 完成并兜底日志

        BackgroundTasks 会以协程方式调度 async 函数，但需保证异常不外泄。
        """
        try:
            await self._parse_and_index(resume_id, file_bytes, file_id, file_name, overwrite)
        except Exception as e:
            logger.exception(f"[BackgroundTasks] 简历 {resume_id} 解析异常: {e}")

    async def _parse_and_index(self, resume_id: str, file_bytes: bytes, file_id: str, file_name: str, overwrite: bool):
        """解析全链路：提取文本 → LLM 结构化 → 去重 → 脱敏 → 父子块 → 入库

        异常会被捕获并标记 parse_status=failed，不会向上抛出。
        向量索引失败仅 warning，不影响结构化解析结果（参照 HRCopilot 设计）。
        """
        try:
            # 1. 文本提取（两级降级）
            text = self._extract_text(file_bytes, file_name)
            # 2. LLM 结构化提取
            structured = await self._llm_extract(text)
            structured = structured or {}
            # 3. 正则兜底补充：LLM 可能漏提取 phone/email/name/skills，用正则从原文补充
            regex_result = _regex_extract_fallback(text, file_name)
            # 缺失字段用正则结果补充（仅当 LLM 返回为空时）
            for key in ("name", "phone", "email", "gender"):
                if not structured.get(key) or not str(structured.get(key, "")).strip():
                    if regex_result.get(key):
                        structured[key] = regex_result[key]
            # skills 为空时用正则补充
            skills = structured.get("skills")
            if not isinstance(skills, list) or len(skills) == 0:
                structured["skills"] = regex_result.get("skills", [])
            # projects 为空时用正则补充
            projects = structured.get("projects")
            if not isinstance(projects, list) or len(projects) == 0:
                structured["projects"] = regex_result.get("projects", [])
            # 4. 字段 None 兜底
            for k in ("name", "phone", "email", "gender", "location", "education", "summary", "salary"):
                if structured.get(k) is None:
                    structured[k] = ""
            for k in ("age",):
                if structured.get(k) is None:
                    structured[k] = 0
            # age 为 0 时尝试从文本中提取
            if not structured.get("age"):
                age_from_text = _extract_age_from_text(text)
                if age_from_text:
                    structured["age"] = age_from_text
            for k in ("skills", "work_experience", "education_detail", "projects"):
                if not isinstance(structured.get(k), list):
                    structured[k] = []
            # 5. 自动计算 work_years：从工作经历起止日期推算
            if not structured.get("work_years"):
                computed_years = _compute_work_years(structured.get("work_experience", []))
                if computed_years > 0:
                    structured["work_years"] = computed_years
                else:
                    # 从文本中提取 "X年经验" 模式
                    years_match = re.search(r'(\d+)\s*[年]', text)
                    if years_match:
                        structured["work_years"] = int(years_match.group(1))
                    else:
                        structured["work_years"] = 0
            # 6. 自动计算 education_level：从 education 字符串或 education_detail 推算
            education_level = structured.get("education_level", 0)
            if not education_level:
                education_level = _compute_education_level(
                    structured.get("education", ""),
                    structured.get("education_detail", []),
                )
            structured["education_level"] = education_level
            # 7. 去重
            phone_h = hash_phone(structured.get("phone", ""))
            email_h = hash_email(structured.get("email", ""))
            dedup_checker = self.dedup_checker or DedupChecker(self.resumes_coll)
            existing = await dedup_checker.check(phone_h, email_h)
            if existing:
                if overwrite:
                    # 覆盖模式：删除旧简历的 MinIO 文件 + MongoDB 记录 + Milvus 向量
                    try:
                        old_doc = await self.resumes_coll.find_one(
                            {"resume_id": existing}, {"file_info": 1}
                        )
                        if old_doc and old_doc.get("file_info", {}).get("file_id"):
                            self.minio.delete(old_doc["file_info"]["file_id"])
                        await self.resumes_coll.delete_one({"resume_id": existing})
                        await self.vector_store.delete_by_resume_id(existing)
                        logger.info(f"覆盖模式：已删除旧简历 {existing}")
                    except Exception as del_err:
                        logger.warning(f"覆盖删除旧简历 {existing} 失败: {del_err}")
                else:
                    # 非覆盖：标记当前简历为重复
                    await self.resumes_coll.update_one(
                        {"resume_id": resume_id},
                        update={"$set": {
                            "is_duplicate": True, "duplicate_with": existing,
                            "parse_info": {"parse_status": "completed", "parse_time": datetime.now(timezone.utc).isoformat()},
                        }},
                    )
                    logger.info(f"简历 {resume_id} 重复，关联 {existing}")
                    return
            # 8. 提前解析 salary（Milvus 写入需要 scalar 字段）
            salary = parse_salary(structured.get("salary", "")) or {"min": 0, "max": 0}
            # 9. 父子块切分
            children, parents = split_parent_child(text)
            # 10. BGE-M3 编码 + 11. 写入 Milvus（索引失败仅 warning，不拖垮 parse_status）
            indexed = True
            skills_list = structured.get("skills", []) or []
            skills_text = " ".join(skills_list)
            try:
                dense, sparse = self.embedding.encode([c.content for c in children])
                await self.vector_store.insert(
                    children, dense, sparse, parents, resume_id,
                    work_years=structured.get("work_years", 0),
                    education_level=structured.get("education_level", 1),
                    salary_min=salary.get("min", 0),
                    salary_max=salary.get("max", 0),
                    skills_text=skills_text,
                )
            except Exception as idx_err:
                indexed = False
                logger.warning(f"简历 {resume_id} 向量索引失败（不影响结构化解析）: {idx_err}")
            # 12. 更新 MongoDB 元数据
            now = datetime.now(timezone.utc).isoformat()
            await self.resumes_coll.update_one(
                {"resume_id": resume_id},
                update={"$set": {
                    "basic_info": {
                        "name": structured.get("name", ""),
                        "phone": structured.get("phone", ""),
                        "email": structured.get("email", ""),
                        "phone_masked": mask_phone(structured.get("phone", "")),
                        "email_masked": mask_email(structured.get("email", "")),
                        "phone_hash": phone_h, "email_hash": email_h,
                        "gender": structured.get("gender"), "age": structured.get("age"),
                        "location": structured.get("location"),
                    },
                    "education": structured.get("education", ""), "education_level": structured.get("education_level", 1),
                    "work_years": structured.get("work_years", 0), "skills": structured.get("skills", []),
                    "work_experience": structured.get("work_experience", []),
                    "education_detail": structured.get("education_detail", []),
                    "projects": structured.get("projects", []),
                    "summary": structured.get("summary", ""),
                    "expected_salary": salary,
                    "parse_info": {"parse_status": "completed", "parse_time": now},
                    "indexed": indexed,
                    "updated_at": now,
                }},
            )
            logger.info(f"简历 {resume_id} 解析完成 (indexed={indexed})")
        except Exception as e:
            logger.exception(f"简历 {resume_id} 解析失败: {e}")
            await self.resumes_coll.update_one(
                {"resume_id": resume_id},
                update={"$set": {"parse_info": {"parse_status": "failed", "parse_time": datetime.now(timezone.utc).isoformat()}}},
            )

    def _extract_text(self, file_bytes: bytes, file_name: str) -> str:
        """两级降级：PyMuPDF/python-docx → RapidOCR

        入参:
            file_bytes: 文件字节
            file_name: 文件名（用于判断扩展名）
        出参:
            提取并清洗后的文本
        异常:
            ResumeParseError: 不支持的文件格式
        """
        ext = file_name.split(".")[-1].lower()
        if ext == "pdf":
            try:
                raw = self._extract_pdf_text(file_bytes)
                return _clean_text(raw)
            except Exception as e:
                logger.warning(f"PDF 主提取失败，降级 OCR: {e}")
                return _clean_text(self.ocr.extract_text(file_bytes))
        elif ext == "docx":
            try:
                raw = self._extract_docx_text(file_bytes)
                return _clean_text(raw)
            except Exception as e:
                logger.warning(f"DOCX 主提取失败，降级 OCR: {e}")
                return _clean_text(self.ocr.extract_text(file_bytes))
        elif ext in ("png", "jpg", "jpeg"):
            return _clean_text(self.ocr.extract_text(file_bytes))
        raise ResumeParseError(f"不支持的文件格式: {ext}")

    @staticmethod
    def _extract_pdf_text(file_bytes: bytes) -> str:
        """PDF 文本提取：PyMuPDF get_text("dict") 按块拼接，保留阅读顺序

        入参:
            file_bytes: PDF 字节
        出参:
            原始文本
        """
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        for page in doc:
            # get_text("dict") 按块按行按 span 输出，保持阅读顺序
            page_dict = page.get_text("dict")
            lines = []
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                    line_text = line_text.strip()
                    if line_text:
                        lines.append(line_text)
            pages_text.append("\n".join(lines))
        doc.close()
        return "\n".join(pages_text)

    @staticmethod
    def _extract_docx_text(file_bytes: bytes) -> str:
        """DOCX 文本提取：XML 全量提取（含文本框/页眉页脚）+ 表格内容

        入参:
            file_bytes: DOCX 字节
        出参:
            原始文本
        """
        import io
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(io.BytesIO(file_bytes))
        text_parts: list[str] = []

        # 方法1：XML 递归提取所有 w:t 标签（支持文本框、形状内文本、页眉页脚）
        try:
            def _collect(element):
                for t in element.iter(qn("w:t")):
                    if t.text:
                        text_parts.append(t.text)
                for br in element.iter(qn("w:br")):
                    text_parts.append("\n")

            _collect(doc._body._element)
            for section in doc.sections:
                if section.header and section.header._element is not None:
                    _collect(section.header._element)
                if section.footer and section.footer._element is not None:
                    _collect(section.footer._element)
        except Exception as e:
            logger.warning(f"DOCX XML 提取失败，回退段落+表格: {e}")
            text_parts = []

        # 方法2：段落 + 表格提取（作为补充/回退）
        para_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        table_texts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_texts.append(" | ".join(cells))

        # 如果 XML 提取的文本为空，直接用段落+表格
        xml_text = "".join(text_parts).strip()
        if not xml_text:
            return "\n".join(para_texts + table_texts)

        # XML 提取有内容，合并段落+表格补充（避免漏掉表格数据）
        all_parts = [xml_text]
        if table_texts:
            all_parts.append("\n".join(table_texts))
        return "\n".join(all_parts)

    async def _llm_extract(self, text: str) -> dict:
        """LLM 结构化提取

        入参:
            text: 简历文本
        出参:
            结构化字典，失败返回 {}
        注意:
            DeepSeek/OpenAI 兼容接口需显式传 response_format=json_object 才能稳定返回 JSON，
            否则可能返回 markdown 代码块包裹或全 null 的 JSON。
            解析时需兼容 ```json ... ``` 包裹与首尾多余文本。
        """
        from app.agent.prompts import RESUME_EXTRACT_PROMPT
        # 截断到 12000 字符（DeepSeek 上下文足够，且简历通常不超过此长度）
        prompt = RESUME_EXTRACT_PROMPT.format(text=str(text)[:12000])
        result = await self.llm.chat(
            [
                {"role": "system", "content": "你是专业的简历信息提取助手。你的任务是从简历文本中准确提取所有字段信息。必须：1)所有字段必须填充，缺失值用空字符串/0/空数组，禁止返回null；2)仔细阅读全文，不要遗漏工作经历、教育经历和技能；3)从工作经历的起止日期推算工作年限；4)只返回JSON，不要任何其他文字。"},
                {"role": "user", "content": prompt},
            ],
            response_format={"type": "json_object"},
        )
        return _safe_json_loads(result)

    async def get_detail(self, resume_id: str, current_user: dict = None) -> dict:
        """AC4.1-4.4: 获取简历详情（根据 current_user.role 返回原始/脱敏手机邮箱）

        入参:
            resume_id: 简历 ID
            current_user: 当前用户（含 role 字段），决定返回原始/脱敏手机邮箱
        出参:
            简历详情字典
        异常:
            NotFoundError: 简历不存在
        """
        doc = await self.resumes_coll.find_one({"resume_id": resume_id}, {"_id": 0})
        if not doc:
            raise NotFoundError("简历不存在")
        # 根据 RBAC 决定 basic_info 中的手机邮箱字段
        basic = dict(doc.get("basic_info") or {})
        is_admin = bool(current_user and current_user.get("role") == "admin")
        if is_admin:
            # admin 看原始值，旧文档无 phone/email 字段时兜底用 masked
            basic["phone"] = basic.get("phone") or basic.get("phone_masked", "")
            basic["email"] = basic.get("email") or basic.get("email_masked", "")
            basic.pop("phone_masked", None)
            basic.pop("email_masked", None)
        else:
            # 普通用户/未登录移除原始字段，保留 masked
            basic.pop("phone", None)
            basic.pop("email", None)
        doc["basic_info"] = basic
        # 扁平化 basic_info 顶层字段（前端期望 name/gender/age/location 在顶层）
        for key in ("name", "gender", "age", "location", "phone_masked", "email_masked", "phone", "email"):
            if key not in doc and key in basic:
                doc[key] = basic[key]
        # name 兜底（避免 basic_info 缺失时前端拿到 None）
        if not doc.get("name"):
            doc["name"] = doc.get("name", "") or ""
        return doc

    async def list(self, page: int = 1, page_size: int = 20, keyword: str = None, tag: str = None,
                   is_favorite: bool = None, education_min: int = None, work_years_min: int = None,
                   salary_min: int = None, salary_max: int = None, status: str = None,
                   date_from: str = None, date_to: str = None,
                   sort_by: str = "created_at", sort_order: str = "desc",
                   current_user: dict = None) -> dict:
        """AC3.1-3.8: 分页列表查询（根据 current_user.role 返回原始/脱敏手机邮箱）

        入参:
            page: 页码（从 1 开始）
            page_size: 每页条数
            keyword: 姓名/技能/标签/摘要关键词
            tag: 标签过滤
            is_favorite: 收藏过滤
            education_min: 最低学历
            work_years_min: 最低工作年限
            salary_min / salary_max: 薪资区间
            status: 解析状态过滤
            date_from: 入库开始日期（ISO 格式 YYYY-MM-DD）
            date_to: 入库结束日期（ISO 格式 YYYY-MM-DD）
            sort_by: 排序字段（created_at/work_years/education_level）
            sort_order: 排序方向（asc/desc）
            current_user: 当前用户（含 role 字段），决定返回原始/脱敏手机邮箱
        出参:
            {"list", "total", "page", "page_size", "total_pages"}
            list 中每项已将 basic_info/parse_info 扁平化到顶层，便于前端卡片直接消费
        """
        query = {}
        if keyword:
            query["$or"] = [
                {"basic_info.name": {"$regex": keyword, "$options": "i"}},
                {"skills": {"$elemMatch": {"$regex": keyword, "$options": "i"}}},
                {"tags": {"$elemMatch": {"$regex": keyword, "$options": "i"}}},
                {"summary": {"$regex": keyword, "$options": "i"}},
            ]
        if tag:
            query["tags"] = tag
        if is_favorite is not None:
            query["is_favorite"] = is_favorite
        if education_min is not None:
            query["education_level"] = {"$gte": education_min}
        if work_years_min is not None:
            query["work_years"] = {"$gte": work_years_min}
        if salary_max is not None:
            query["expected_salary.min"] = {"$lte": salary_max}
        if salary_min is not None:
            query["expected_salary.max"] = {"$gte": salary_min}
        if status:
            query["parse_info.parse_status"] = status
        # 日期范围筛选（自动交换倒序日期）
        if date_from and date_to and date_from > date_to:
            date_from, date_to = date_to, date_from
        if date_from:
            query["created_at"] = {"$gte": date_from}
        if date_to:
            query.setdefault("created_at", {})["$lte"] = date_to + "T23:59:59"
        total = await self.resumes_coll.count_documents(query)
        skip = (page - 1) * page_size
        # sort_by 白名单校验，非法值兜底为 created_at
        ALLOWED_SORT_FIELDS = {"created_at", "work_years", "education_level"}
        safe_sort_by = sort_by if sort_by in ALLOWED_SORT_FIELDS else "created_at"
        safe_sort_order = 1 if sort_order == "asc" else -1
        cursor = self.resumes_coll.find(query, {"_id": 0}).skip(skip).limit(page_size).sort(safe_sort_by, safe_sort_order)
        raw_items = await cursor.to_list(length=page_size)
        # 扁平化：前端 ResumeListItem 期望 name/gender/age/location/parse_status 在顶层
        items = [self._flatten_for_list(it, current_user) for it in raw_items]
        return {
            "list": items, "total": total, "page": page, "page_size": page_size,
            "total_pages": (total + page_size - 1) // page_size,
        }

    @staticmethod
    def _flatten_for_list(doc: dict, current_user: dict = None) -> dict:
        """将 MongoDB 文档的 basic_info/parse_info 扁平化到顶层，根据 current_user 角色返回原始/脱敏手机邮箱

        入参:
            doc: MongoDB 原始文档
            current_user: 当前用户信息（含 role 字段），决定返回 phone/email 或 masked 字段
        出参:
            扁平化后的字典，顶层包含 name/gender/age/location/parse_status 等字段
            admin 用户看到 phone/email（旧文档无原始值时兜底用 masked）
            普通用户/未登录看到 phone_masked/email_masked
        """
        if not doc:
            return doc
        basic = doc.get("basic_info") or {}
        parse = doc.get("parse_info") or {}
        flat = dict(doc)
        # 顶层补充前端期望的扁平字段（若文档本身已有顶层字段则不覆盖）
        for k in ("name", "gender", "age", "location"):
            if k not in flat and k in basic:
                flat[k] = basic[k]
        if "name" not in flat:
            flat["name"] = basic.get("name", "")
        if "parse_status" not in flat:
            flat["parse_status"] = parse.get("parse_status", "pending")
        # 根据 RBAC 决定返回原始值还是 masked 值
        is_admin = bool(current_user and current_user.get("role") == "admin")
        if is_admin:
            # admin 看原始值，旧文档无 phone/email 字段时兜底用 masked
            flat["phone"] = basic.get("phone") or basic.get("phone_masked", "")
            flat["email"] = basic.get("email") or basic.get("email_masked", "")
            flat.pop("phone_masked", None)
            flat.pop("email_masked", None)
        else:
            # 普通用户/未登录看 masked 值，移除原始字段（防止泄露）
            flat["phone_masked"] = basic.get("phone_masked", "")
            flat["email_masked"] = basic.get("email_masked", "")
            flat.pop("phone", None)
            flat.pop("email", None)
        return flat

    async def delete(self, resume_id: str) -> None:
        """AC6.1-6.4: 清理 MinIO/MongoDB/Milvus

        入参:
            resume_id: 简历 ID
        """
        doc = await self.resumes_coll.find_one({"resume_id": resume_id}, {"file_info": 1})
        if doc and doc.get("file_info", {}).get("file_id"):
            self.minio.delete(doc["file_info"]["file_id"])
        await self.resumes_coll.delete_one({"resume_id": resume_id})
        await self.vector_store.delete_by_resume_id(resume_id)

    async def get_preview_url(self, resume_id: str) -> dict:
        """AC5.1-5.2: 生成文件预签名 URL

        入参:
            resume_id: 简历 ID
        出参:
            {"preview_url", "file_type", "expires_in"}
        异常:
            NotFoundError: 简历不存在
        """
        doc = await self.resumes_coll.find_one({"resume_id": resume_id}, {"file_info": 1})
        if not doc:
            raise NotFoundError("简历不存在")
        url = self.minio.presigned_url(doc["file_info"]["file_id"])
        return {"preview_url": url, "file_type": doc["file_info"]["file_type"], "expires_in": 3600}


def _safe_json_loads(result: str) -> dict:
    """robust JSON 解析：去除 markdown 代码块包裹后 json.loads，失败返回 {}

    入参:
        result: LLM 返回的原始文本
    出参:
        解析后的字典；失败返回 {}
    """
    if not result:
        return {}
    text = result.strip()
    # 去除 ```json ... ``` 或 ``` ... ``` 包裹
    if text.startswith("```"):
        # 去掉首行 ```json / ```
        lines = text.split("\n", 1)
        if len(lines) > 1:
            text = lines[1]
        # 去掉末尾 ```
        if text.endswith("```"):
            text = text[: text.rfind("```")]
        text = text.strip()
    try:
        return json.loads(text)
    except Exception:
        logger.error(f"LLM 返回非 JSON: {str(result)[:200]}")
        return {}


def _clean_text(raw: str) -> str:
    """文本清洗：压缩多余空白、去除控制字符、统一换行

    入参:
        raw: 原始提取文本
    出参:
        清洗后的文本
    """
    if not raw:
        return ""
    if not isinstance(raw, str):
        raw = str(raw)
    # 去除零宽字符和控制字符（保留换行 \n 和制表符 \t）
    cleaned = re.sub(r'[\u200b-\u200f\u202a-\u202e\ufeff\x00-\x08\x0b\x0c\x0e-\x1f]', '', raw)
    # 合并连续空行（3个以上换行 → 2个）
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    # 合并行内连续空格（2个以上 → 1个）
    cleaned = re.sub(r'[ \t]{2,}', ' ', cleaned)
    return cleaned.strip()


def _regex_extract_fallback(raw_text: str, file_name: str = "") -> dict:
    """正则降级提取：当 LLM 漏提取关键字段时，用正则从原文补充

    入参:
        raw_text: 简历原始文本
        file_name: 文件名（用于从文件名猜测姓名）
    出参:
        包含 name/phone/email/gender/skills 的字典
    """
    result = {"name": "", "phone": "", "email": "", "gender": "", "skills": []}

    if not isinstance(raw_text, str):
        raw_text = str(raw_text) if raw_text else ""
    if not raw_text:
        return result

    # 手机号：中国大陆 11 位手机号（1开头，第二位 3-9）
    phone_match = re.search(r'(?<!\d)(1[3-9]\d{9})(?!\d)', raw_text)
    if phone_match:
        result["phone"] = phone_match.group(1)

    # 邮箱
    email_match = re.search(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}', raw_text)
    if email_match:
        result["email"] = email_match.group(0)

    # 性别
    if re.search(r'(性别\s*[:：]?\s*男|男\s*[|｜/]\s*\d+\s*岁|\b男\b)', raw_text):
        result["gender"] = "男"
    elif re.search(r'(性别\s*[:：]?\s*女|女\s*[|｜/]\s*\d+\s*岁|\b女\b)', raw_text):
        result["gender"] = "女"

    # 姓名：优先从文件名提取（常见格式：张三-简历.docx, 张三_Java开发.docx）
    name_from_file = ""
    if file_name:
        base = file_name.rsplit('.', 1)[0] if '.' in file_name else file_name
        # 用 - _ — 分割取第一段，去掉"简历""个人简历"等关键词
        first_part = re.split(r'[-_—\s]', base)[0].strip()
        if 1 < len(first_part) <= 10 and not re.search(r'\d', first_part) and first_part not in ("简历", "个人简历", "求职简历"):
            name_from_file = first_part

    # 文件名没提取到，从文本前 500 字中找最像姓名的行
    if not name_from_file:
        head = raw_text[:500]
        lines = [line.strip() for line in head.split('\n') if line.strip()][:10]
        for line in lines:
            # 排除含数字、@、简历关键词、邮箱电话标签的行
            if 1 < len(line) <= 10 and not re.search(r'\d|@|简历|求职|个人|信息|电话|邮箱|手机|学历|专业|学校|公司', line):
                name_from_file = line
                break
    result["name"] = name_from_file

    # 技能关键词匹配
    common_skills = [
        'Python', 'Java', 'Go', 'Golang', 'C++', 'C#', 'JavaScript', 'TypeScript',
        'Vue', 'Vue.js', 'React', 'React.js', 'Angular', 'Node.js',
        'Spring', 'Spring Boot', 'Spring Cloud', 'Django', 'Flask', 'FastAPI',
        'MySQL', 'PostgreSQL', 'Oracle', 'SQL Server', 'Redis', 'MongoDB', 'Elasticsearch',
        'Linux', 'Docker', 'Kubernetes', 'K8s', 'Git', 'GitLab', 'Jenkins',
        'Nginx', 'RabbitMQ', 'Kafka', 'Zookeeper',
        'HTML', 'CSS', 'Sass', 'Less', 'Webpack', 'Vite',
        '机器学习', '深度学习', 'NLP', '自然语言处理', 'CV', '计算机视觉', 'TensorFlow', 'PyTorch',
        '微服务', '分布式', '高并发', 'Hadoop', 'Spark', 'Flink', 'Hive',
        'MyBatis', 'MyBatis-Plus', 'Hibernate', 'JPA',
        '数据结构', '算法', '网络编程', '多线程',
        '产品经理', '项目管理', 'PMP', '敏捷开发', 'Scrum',
        'Excel', 'Word', 'PPT', 'PowerPoint', '数据分析', 'SQL',
        'Android', 'iOS', 'Swift', 'Kotlin', 'Flutter', 'React Native',
        '运维', 'DevOps', 'CI/CD', 'AWS', '阿里云', '腾讯云',
    ]
    found_skills = []
    text_lower = raw_text.lower()
    for skill in common_skills:
        if skill.lower() in text_lower:
            found_skills.append(skill)
    result["skills"] = found_skills[:20]

    # 项目经历兜底提取：从文本中识别"项目名称:"/"项目经历"等模式
    result["projects"] = _extract_projects_from_text(raw_text)

    return result


def _extract_projects_from_text(raw_text: str) -> list:
    """从简历文本中正则提取项目经历

    入参:
        raw_text: 简历原始文本
    出参:
        项目列表，每项含 name/role/description
    """
    if not raw_text:
        return []
    projects = []
    # 匹配 "项目名称：XXX" / "项目：XXX" / "项目经历：XXX" 等
    project_pattern = re.compile(
        r'(?:项目名称|项目|项目经历|项目经验)\s*[:：]\s*(.+?)(?=\n(?:项目|工作|教育|技能|自我|个人|证书|语言|兴趣爱好|$))',
        re.DOTALL
    )
    for m in project_pattern.finditer(raw_text):
        block = m.group(1).strip()
        if not block:
            continue
        lines = [l.strip() for l in block.split('\n') if l.strip()]
        name = lines[0] if lines else ""
        description = "\n".join(lines[1:]) if len(lines) > 1 else ""
        projects.append({
            "name": name[:100],
            "role": "",
            "description": description[:500],
        })
    return projects[:10]  # 最多 10 个项目


def _extract_age_from_text(raw_text: str) -> int:
    """从文本中提取年龄

    入参:
        raw_text: 简历文本
    出参:
        年龄整数；提取不到返回 0
    """
    # 模式1: "XX岁"
    age_match = re.search(r'(\d{2})\s*岁', raw_text)
    if age_match:
        age = int(age_match.group(1))
        if 16 <= age <= 70:
            return age

    # 模式2: "出生于19XX年" 或 "19XX年出生" → 推算年龄
    birth_match = re.search(r'(19[89]\d|20[01]\d)\s*年?\s*出生', raw_text)
    if birth_match:
        birth_year = int(birth_match.group(1))
        current_year = datetime.now().year
        age = current_year - birth_year
        if 16 <= age <= 70:
            return age

    return 0


def _compute_work_years(work_experience: list) -> int:
    """从工作经历列表推算总工作年限

    入参:
        work_experience: LLM 提取的工作经历列表，每项含 start_date/end_date
    出参:
        工作年限整数；无法推算返回 0
    """
    if not work_experience or not isinstance(work_experience, list):
        return 0
    total_months = 0
    now = datetime.now(timezone.utc)
    for exp in work_experience:
        if not isinstance(exp, dict):
            continue
        start_str = str(exp.get("start_date", "") or "").strip()
        end_str = str(exp.get("end_date", "") or "").strip()
        if not start_str:
            continue
        try:
            # 解析 start_date: YYYY-MM
            start = datetime.strptime(start_str[:7], "%Y-%m")
            # 解析 end_date: "至今" 或为空则用当前时间；否则解析 YYYY-MM
            if not end_str or end_str in ("至今", "现在", "在职", "present", "null", "None"):
                end = now
            else:
                end = datetime.strptime(end_str[:7], "%Y-%m")
            months = (end.year - start.year) * 12 + (end.month - start.month)
            if months > 0:
                total_months += months
        except (ValueError, TypeError):
            continue
    return max(total_months // 12, 0)


def _compute_education_level(education_str: str, education_detail: list) -> int:
    """从学历字符串或教育经历列表推算 education_level

    入参:
        education_str: LLM 返回的最高学历字符串
        education_detail: LLM 返回的教育经历列表
    出参:
        0=专科/高中及以下, 1=本科, 2=硕士, 3=博士；默认本科(1)
    """
    # 优先级：博士 > 硕士 > 本科 > 专科
    level_keywords = [
        (3, ["博士", "phd", "Ph.D", "doctor"]),
        (2, ["硕士", "研究生", "master", "MBA", "mba"]),
        (1, ["本科", "学士", "bachelor"]),
        (0, ["专科", "大专", "高职", "高中", "中专"]),
    ]

    text = education_str.lower()
    # 也检查 education_detail 中的 degree 字段
    if isinstance(education_detail, list):
        for edu in education_detail:
            if isinstance(edu, dict):
                deg = str(edu.get("degree", "") or "").lower()
                text += " " + deg

    for level, keywords in level_keywords:
        for kw in keywords:
            if kw.lower() in text:
                return level

    # 默认返回本科(1)作为合理默认
    return 1
